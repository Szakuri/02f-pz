from flask import ( # type: ignore
    Flask,
    render_template,
    request,
    jsonify,
    redirect,
    url_for,
    flash,
    session,
)

from flask_sqlalchemy import SQLAlchemy 
from flask_migrate import Migrate
from werkzeug.security import generate_password_hash, check_password_hash 
from pytesseract import image_to_string 
from pdf2image import convert_from_path
from PIL import Image
from docx import Document 
import os


db = SQLAlchemy()
migrate = Migrate()


# Funkcja do tworzenia aplikacji
def create_app():
    app = Flask(__name__)
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///database.db"
    app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
    app.config["UPLOAD_FOLDER"] = os.path.join(
        os.path.abspath(os.path.dirname(__file__)), "uploads"
    )
    app.secret_key = os.urandom(24)
    app.secret_key = os.getenv("SECRET_KEY", "domyslny_klucz")

    # Inicjalizacja bazy danych z aplikacją
    db.init_app(app)
<<<<<<< Updated upstream
    migrate = Migrate(app, db)

    # Importowanie modeli i tworzenie globalnych stanowisk
    with app.app_context():
        from models import Position, Keyword, Candidate, User

        db.create_all()
        create_global_positions()
=======
    migrate.init_app(app, db)
        
    # Importowanie modeli
    with app.app_context():
        from models import Position, Keyword, Candidate, User
        db.create_all() 
        create_default_positions()
>>>>>>> Stashed changes

    # Zdefiniowanie endpointów
    @app.route("/")
    def home():
        if "user_id" not in session:
            flash("Musisz się zalogować, aby zobaczyć tę stronę.")
            return redirect(url_for("login"))
        user = User.query.get(session["user_id"])
        return render_template("home.html", username=user.username)

    @app.route("/upload")
    def upload():
        if "user_id" not in session:
            flash("Musisz się zalogować!")
            return redirect(url_for("login"))

        user_id = session["user_id"]
<<<<<<< Updated upstream
        user_positions = Position.query.filter_by(user_id=user_id).all()
        global_positions = Position.query.filter_by(is_global=True).all()
        return render_template(
            "upload.html",
            user_positions=user_positions,
            global_positions=global_positions,
=======

        # Pobierz globalne i użytkownika stanowiska
        global_positions = Position.query.filter_by(is_default=True).all()
        user_positions = Position.query.filter_by(user_id=user_id).all()

        return render_template(
            "upload.html",
            global_positions=global_positions,
            user_positions=user_positions
>>>>>>> Stashed changes
        )

    @app.route("/positions", methods=["POST"])
    def add_position():
        data = request.json
        position = Position(title=data["title"])
        db.session.add(position)
        db.session.commit()

        for word in data["keywords"]:
            keyword = Keyword(word=word, position_id=position.id)
            db.session.add(keyword)

        db.session.commit()
        return jsonify({"message": "Stanowisko zostało pomyślnie dodane!"}), 201

    # Endpoint: Analiza CV
    @app.route("/analyze_cv", methods=["POST"])
    def analyze_cv():
        try:
            # Pobierz dane z formularza
            name = request.form["name"]
            position_id = int(request.form["position_id"])
            file = request.files["file"]

            # Zapis pliku
            file_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
            file.save(file_path)

            # Pobieranie słów kluczowych dla stanowiska
            keywords = Keyword.query.filter_by(position_id=position_id).all()

            # Odczytanie tekstu z pliku za pomocą OCR
            extracted_text = ""
            if file.filename.lower().endswith((".png", ".jpg", ".jpeg")):
                extracted_text = image_to_string(file_path)
            elif file.filename.lower().endswith(".pdf"):
                pages = convert_from_path(file_path)
                extracted_text = " ".join(image_to_string(page) for page in pages)
<<<<<<< Updated upstream
            elif file.filename.lower().endswith(".docx"):
                doc = Document(file_path)
                extracted_text = " ".join(
                    paragraph.text for paragraph in doc.paragraphs
                )
=======
            except Exception as e:
                flash(f"Błąd podczas wyodrębniania tekstu z PDF: {str(e)}")
                return redirect(url_for("upload"))
>>>>>>> Stashed changes

            # Normalizacja tekstu
            normalized_text = extracted_text.lower()

            # Analiza słów kluczowych
            results = {}
            total_score = 0
            for keyword in keywords:
                count = normalized_text.count(keyword.word.lower())
                points = count * keyword.weight  # Uwzględnienie wagi
                results[keyword.word] = {"count": count, "weight": keyword.weight, "points": points}

<<<<<<< Updated upstream
            total_score = sum(results.values())

            # Renderowanie szablonu HTML z wynikami
            return render_template(
                "analyze_results.html",
                name=name,
                results=results,
                total_score=total_score,
            )
=======
            total_score += points

            # Zapis kandydata do bazy danych z punktami
            candidate = Candidate(
                name=name,
                cv_text=extracted_text,
                position_id=position_id,
                points=total_score  # Zapis liczby punktów
            )
            db.session.add(candidate)
            db.session.commit()

            # Renderowanie wyników
            return render_template("results.html", name=name, results=results, total_score=total_score)
>>>>>>> Stashed changes

        except Exception as e:
            flash(f"Wystąpił błąd: {str(e)}")
            return redirect(url_for("upload"))

    @app.route("/add_position", methods=["GET", "POST"])
    def add_position_form():
        if "user_id" not in session:
            flash("Musisz się zalogować!")
            return redirect(url_for("login"))

        if request.method == "POST":
            title = request.form["title"]
            keywords = request.form["keywords"].split(",")

            position = Position(title=title, user_id=session["user_id"])
            db.session.add(position)
            db.session.commit()

            for keyword in keywords:
                kw = Keyword(word=keyword.strip(), position_id=position.id)
                db.session.add(kw)

            db.session.commit()

            flash("Stanowisko zostało dodane pomyślnie!")
            return redirect(url_for("home"))

        return render_template("add_position.html")

    @app.route("/register", methods=["GET", "POST"])
    def register():
        if request.method == "POST":
            username = request.form["username"]
            email = request.form["email"]
            password = request.form["password"]

            # Sprawdź, czy użytkownik istnieje
            if User.query.filter_by(username=username).first():
                flash("Nazwa użytkownika jest już zajęta!")
                return redirect(url_for("register"))

            # Utwórz użytkownika
            new_user = User(username=username, email=email)
            new_user.set_password(password)
            db.session.add(new_user)
            db.session.commit()

            flash("Rejestracja zakończona pomyślnie! Teraz możesz się zalogować.")
            return redirect(url_for("login"))

        return render_template("register.html")

    @app.route("/login", methods=["GET", "POST"])
    def login():
        if request.method == "POST":
            username = request.form["username"]
            password = request.form["password"]

            user = User.query.filter_by(username=username).first()

            if user and user.check_password(password):
                session["user_id"] = user.id  # Ustawiamy ID użytkownika w sesji
                flash("Zalogowano pomyślnie!")
                return redirect(url_for("home"))  # Przekierowanie na stronę główną
            else:
                flash("Nieprawidłowa nazwa użytkownika lub hasło.", "error")
                return redirect(url_for("login"))  # Powrót na stronę logowania

        return render_template("login.html")

    @app.route("/logout")
    def logout():
        session.pop("user_id", None)
        flash("Wylogowano pomyślnie!")
        return redirect(url_for("login"))

    @app.route("/view_positions")
    def view_positions():
        if "user_id" not in session:
            flash("Musisz się zalogować!")
            return redirect(url_for("login"))

        user_id = session["user_id"]
        positions = Position.query.filter_by(user_id=user_id).all()
        return render_template("view_positions.html", positions=positions)
    

    @app.route("/ranking", methods=["GET", "POST"])
    def ranking():
        try:
            # Pobierz listę stanowisk
            positions = Position.query.all()

            # Domyślne stanowisko (pierwsze w bazie lub wybrane w formularzu)
            position_id = request.args.get("position_id", type=int) or positions[0].id

            # Pobierz parametr liczby wyników z URL (domyślnie 20)
            limit = request.args.get("limit", default=20, type=int)

            # Walidacja limitu
            if limit < 1:
                limit = 1
            elif limit > 50:
                limit = 50

            # Pobierz stanowisko
            position = Position.query.get_or_404(position_id)

            # Pobierz kandydatów dla stanowiska i posortuj ich według punktów malejąco
            candidates = (
                Candidate.query.filter_by(position_id=position_id)
                .order_by(Candidate.points.desc())
                .limit(limit)
                .all()
            )

            # Renderowanie rankingu
            return render_template(
                "ranking.html",
                positions=positions,
                position=position,
                candidates=candidates,
                limit=limit,
            )

        except Exception as e:
            flash(f"Wystąpił błąd: {str(e)}")
            return redirect(url_for("home"))
        
    @app.route("/edit_position/<int:position_id>", methods=["GET", "POST"])
    def edit_position(position_id):
        position = Position.query.get_or_404(position_id)

        # Sprawdzenie, czy stanowisko jest globalne
        if position.is_default:
            flash("Nie można edytować globalnych stanowisk!")
            return redirect(url_for("view_positions"))

        if request.method == "POST":
            # Zaktualizuj tytuł stanowiska
            title = request.form.get("title")
            position.title = title

            # Pobierz z formularza keyword_ids i weights jako oddzielne listy
            keyword_ids = request.form.getlist("keyword_ids")
            weights = request.form.getlist("weights")

            # Zaktualizuj każdą wagę słowa kluczowego
            for keyword_id, weight in zip(keyword_ids, weights):
                keyword = Keyword.query.get(int(keyword_id))
                if keyword:
                    keyword.weight = int(weight)

            # Zapisz zmiany w bazie danych
            db.session.commit()
            flash("Stanowisko zostało zaktualizowane!")
            return redirect(url_for("view_positions"))

        # Pobierz słowa kluczowe dla stanowiska
        keywords = Keyword.query.filter_by(position_id=position.id).all()

        return render_template("edit_position.html", position=position, keywords=keywords)





    @app.route("/manage_global_positions", methods=["GET", "POST"])
    def manage_global_positions():
        if request.method == "POST":
            title = request.form["title"]
            keywords = request.form["keywords"].split(",")

            existing_position = Position.query.filter_by(
                title=title, is_global=True
            ).first()
            if existing_position:
                flash("Takie stanowisko już istnieje!")
                return redirect(url_for("manage_global_positions"))

            new_position = Position(title=title, is_global=True)
            db.session.add(new_position)
            db.session.commit()

            for keyword in keywords:
                new_keyword = Keyword(word=keyword.strip(), position_id=new_position.id)
                db.session.add(new_keyword)

            db.session.commit()
            flash("Stanowisko globalne zostało dodane!")
            return redirect(url_for("manage_global_positions"))

        global_positions = Position.query.filter_by(is_global=True).all()
        return render_template(
            "manage_global_positions.html", global_positions=global_positions
        )

    @app.route("/upload_multiple", methods=["POST"])
    def upload_multiple():
        try:
            name = request.form["name"]
            position_id = int(request.form["position_id"])
            uploaded_files = request.files.getlist("files")

            if not uploaded_files:
                return jsonify({"error": "Nie przesłano żadnych plików"}), 400

            keywords = [
                kw.word for kw in Keyword.query.filter_by(position_id=position_id).all()
            ]

            # Lista na wyniki analizy dla każdego pliku
            results = []

            for file in uploaded_files:
                file_path = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
                file.save(file_path)

                extracted_text = ""
                if file.filename.lower().endswith((".png", ".jpg", ".jpeg")):
                    extracted_text = image_to_string(Image.open(file_path))
                elif file.filename.lower().endswith(".pdf"):
                    pages = convert_from_path(file_path)
                    extracted_text = " ".join(image_to_string(page) for page in pages)
                elif file.filename.lower().endswith(".docx"):
                    doc = Document(file_path)
                    extracted_text = " ".join(
                        paragraph.text for paragraph in doc.paragraphs
                    )

                normalized_text = extracted_text.lower()

                # Analiza słów kluczowych dla bieżącego pliku
                analysis = {
                    keyword: normalized_text.count(keyword.lower())
                    for keyword in keywords
                }
                total_score = sum(analysis.values())

                # Dodaj analizę do wyników
                results.append(
                    {
                        "name": file.filename,
                        "analysis": analysis,
                        "total_score": total_score,
                    }
                )

            # Przekaż wyniki do szablonu HTML
            return render_template("analyze_results.html", results=results)

        except Exception as e:
            return jsonify({"error": str(e)}), 500

    return app

def create_default_positions():
    from models import Position, Keyword

    Position.query.filter_by(is_default=True).delete()
    db.session.commit()

    # Struktura z wagami dla słów kluczowych
    default_positions = [
        {
        "title": "Programista",
        "keywords": [
            {"word": "programowanie", "weight": 5}, {"word": "Python", "weight": 5},
            {"word": "Java", "weight": 5}, {"word": "JavaScript", "weight": 5},
            {"word": "C#", "weight": 4}, {"word": "SQL", "weight": 4},
            {"word": "bazy danych", "weight": 4}, {"word": "API", "weight": 4},
            {"word": "framework", "weight": 3}, {"word": "debugowanie", "weight": 3},
            {"word": "kodowanie", "weight": 3}, {"word": "systemy operacyjne", "weight": 2},
            {"word": "integracja", "weight": 2}, {"word": "automatyzacja", "weight": 2},
            {"word": "scrum", "weight": 3}, {"word": "agile", "weight": 3},
            {"word": "testing", "weight": 3}, {"word": "backend", "weight": 3},
            {"word": "frontend", "weight": 3}, {"word": "narzędzia CI/CD", "weight": 4},
            {"word": "Git", "weight": 4}, {"word": "Docker", "weight": 4},
            {"word": "Linux", "weight": 5}, {"word": "Windows", "weight": 4},
            {"word": "architektura systemów", "weight": 5}, {"word": "analityka", "weight": 4},
            {"word": "optymalizacja", "weight": 3}, {"word": "REST", "weight": 4},
            {"word": "SOAP", "weight": 3}, {"word": "mikroserwisy", "weight": 4},
            {"word": "bezpieczeństwo IT", "weight": 5}, {"word": "refaktoryzacja", "weight": 4},
            {"word": "chmura", "weight": 5}, {"word": "Azure", "weight": 4},
            {"word": "AWS", "weight": 4}, {"word": "GCP", "weight": 4}, {"word": "DevOps", "weight": 5}, 
            {"word": "B1", "weight": 1}, {"word": "B2", "weight": 2},
            {"word": "C1", "weight": 3}, {"word": "C2", "weight": 4}, {"word": "English", "weight": 1}, 
            {"word": "Angielski", "weight": 1}
        ],
    },
    {
        "title": "Analityk Danych",
        "keywords": [
            {"word": "dane", "weight": 5}, {"word": "Excel", "weight": 5},
            {"word": "SQL", "weight": 5}, {"word": "analizy", "weight": 4},
            {"word": "wizualizacja", "weight": 4}, {"word": "Tableau", "weight": 4},
            {"word": "Power BI", "weight": 4}, {"word": "statystyka", "weight": 5},
            {"word": "Python", "weight": 5}, {"word": "R", "weight": 4},
            {"word": "modele predykcyjne", "weight": 5}, {"word": "algorytmy", "weight": 5},
            {"word": "analiza danych", "weight": 5}, {"word": "przetwarzanie danych", "weight": 4},
            {"word": "ETL", "weight": 4}, {"word": "dashboard", "weight": 4},
            {"word": "raporty", "weight": 4}, {"word": "Big Data", "weight": 5},
            {"word": "Hadoop", "weight": 5}, {"word": "Spark", "weight": 5},
            {"word": "analityka biznesowa", "weight": 4}, {"word": "insighty", "weight": 4},
            {"word": "dane sprzedażowe", "weight": 4}, {"word": "prognozowanie", "weight": 5},
            {"word": "optymalizacja", "weight": 4}, {"word": "bazy danych", "weight": 4},
            {"word": "machine learning", "weight": 5}, {"word": "analiza rynku", "weight": 4},
            {"word": "czyszczenie danych", "weight": 3}, {"word": "Google Analytics", "weight": 4},
            {"word": "KPIs", "weight": 4}, {"word": "modelowanie danych", "weight": 5}, 
            {"word": "B1", "weight": 1}, {"word": "B2", "weight": 2},
            {"word": "C1", "weight": 3}, {"word": "C2", "weight": 4}, {"word": "English", "weight": 1}, 
            {"word": "Angielski", "weight": 1}
        ],
    },
    {
        "title": "Kierownik Projektu",
        "keywords": [
            {"word": "zarządzanie projektami", "weight": 5}, {"word": "harmonogramy", "weight": 4},
            {"word": "budżet", "weight": 5}, {"word": "agile", "weight": 4},
            {"word": "scrum", "weight": 4}, {"word": "waterfall", "weight": 3},
            {"word": "zarządzanie ryzykiem", "weight": 5}, {"word": "priorytetyzacja", "weight": 4},
            {"word": "planowanie", "weight": 4}, {"word": "komunikacja", "weight": 5},
            {"word": "zarządzanie zespołem", "weight": 5}, {"word": "stakeholderzy", "weight": 4},
            {"word": "analiza wymagań", "weight": 4}, {"word": "cele projektu", "weight": 5},
            {"word": "raportowanie", "weight": 4}, {"word": "monitoring", "weight": 4},
            {"word": "zarządzanie czasem", "weight": 4}, {"word": "zarządzanie zmianą", "weight": 5},
            {"word": "leadership", "weight": 5}, {"word": "efektywność", "weight": 4},
            {"word": "metodologia", "weight": 4}, {"word": "stand-up", "weight": 3},
            {"word": "roadmap", "weight": 3}, {"word": "narzędzia PM", "weight": 4},
            {"word": "Jira", "weight": 5}, {"word": "Asana", "weight": 4},
            {"word": "MS Project", "weight": 4}, {"word": "zarządzanie zadaniami", "weight": 4}, 
            {"word": "B1", "weight": 1}, {"word": "B2", "weight": 2},
            {"word": "C1", "weight": 3}, {"word": "C2", "weight": 4}, {"word": "English", "weight": 1}, 
            {"word": "Angielski", "weight": 1}
        ],
    },

    {
        "title": "Grafik",
        "keywords": [
            {"word": "projektowanie graficzne", "weight": 5}, {"word": "Adobe Photoshop", "weight": 5},
            {"word": "Illustrator", "weight": 4}, {"word": "InDesign", "weight": 4},
            {"word": "Canva", "weight": 3}, {"word": "projektowanie logo", "weight": 4},
            {"word": "typografia", "weight": 3}, {"word": "paleta kolorów", "weight": 3},
            {"word": "branding", "weight": 4}, {"word": "layout", "weight": 3},
            {"word": "infografiki", "weight": 4}, {"word": "kreatywność", "weight": 5},
            {"word": "ilustracje", "weight": 4}, {"word": "UX/UI", "weight": 5},
            {"word": "projektowanie interfejsów", "weight": 5}, {"word": "responsywność", "weight": 4},
            {"word": "druk", "weight": 3}, {"word": "składy", "weight": 2},
            {"word": "grafika wektorowa", "weight": 4}, {"word": "animacje", "weight": 5},
            {"word": "mockupy", "weight": 3}, {"word": "wizualizacje", "weight": 3},
            {"word": "wireframes", "weight": 4}, {"word": "projektowanie banerów", "weight": 3},
            {"word": "grafika 3D", "weight": 5}, {"word": "narzędzia graficzne", "weight": 4}, 
            {"word": "B1", "weight": 1}, {"word": "B2", "weight": 2},
            {"word": "C1", "weight": 3}, {"word": "C2", "weight": 4}, {"word": "English", "weight": 1}, 
            {"word": "Angielski", "weight": 1}
        ],
    },
    {
        "title": "Copywriter",
        "keywords": [
            {"word": "SEO", "weight": 5}, {"word": "optymalizacja treści", "weight": 5},
            {"word": "copywriting", "weight": 5}, {"word": "marketing treści", "weight": 4},
            {"word": "artykuły", "weight": 4}, {"word": "blogi", "weight": 4},
            {"word": "storytelling", "weight": 5}, {"word": "kreatywne pisanie", "weight": 5},
            {"word": "redagowanie", "weight": 4}, {"word": "social media", "weight": 4},
            {"word": "CTA", "weight": 3}, {"word": "reklama", "weight": 4},
            {"word": "edytowanie", "weight": 4}, {"word": "język", "weight": 5},
            {"word": "treść", "weight": 4}, {"word": "analiza konkurencji", "weight": 4},
            {"word": "Google Analytics", "weight": 4}, {"word": "targetowanie", "weight": 3},
            {"word": "kampanie marketingowe", "weight": 4}, {"word": "e-mail marketing", "weight": 4},
            {"word": "UX writing", "weight": 4}, {"word": "optymalizacja słów kluczowych", "weight": 5},
            {"word": "analiza odbiorców", "weight": 4}, 
            {"word": "B1", "weight": 1}, {"word": "B2", "weight": 2},
            {"word": "C1", "weight": 3}, {"word": "C2", "weight": 4}, {"word": "English", "weight": 1}, 
            {"word": "Angielski", "weight": 1}
        ],
    },
    {
        "title": "Specjalista ds. Marketingu Cyfrowego",
        "keywords": [
            {"word": "marketing cyfrowy", "weight": 5}, {"word": "SEO", "weight": 5},
            {"word": "SEM", "weight": 5}, {"word": "Google Ads", "weight": 4},
            {"word": "Facebook Ads", "weight": 4}, {"word": "social media", "weight": 4},
            {"word": "kampanie", "weight": 4}, {"word": "reklama", "weight": 4},
            {"word": "targetowanie", "weight": 4}, {"word": "optymalizacja", "weight": 5},
            {"word": "Google Analytics", "weight": 5}, {"word": "email marketing", "weight": 4},
            {"word": "automatyzacja marketingu", "weight": 4}, {"word": "content marketing", "weight": 4},
            {"word": "remarketing", "weight": 4}, {"word": "KPI", "weight": 3},
            {"word": "lead generation", "weight": 5}, {"word": "konwersje", "weight": 4},
            {"word": "budżet reklamowy", "weight": 4}, {"word": "strategia marketingowa", "weight": 5}, 
            {"word": "B1", "weight": 1}, {"word": "B2", "weight": 2},
            {"word": "C1", "weight": 3}, {"word": "C2", "weight": 4}, {"word": "English", "weight": 1}, 
            {"word": "Angielski", "weight": 1}
        ],
    },
    {
        "title": "Administrator Sieci",
        "keywords": [
            {"word": "sieci komputerowe", "weight": 5}, {"word": "protokoły", "weight": 5},
            {"word": "TCP/IP", "weight": 5}, {"word": "DNS", "weight": 4},
            {"word": "DHCP", "weight": 4}, {"word": "routing", "weight": 4},
            {"word": "VPN", "weight": 5}, {"word": "bezpieczeństwo", "weight": 5},
            {"word": "firewalle", "weight": 5}, {"word": "zarządzanie siecią", "weight": 5},
            {"word": "konfiguracja", "weight": 4}, {"word": "LAN", "weight": 4},
            {"word": "WAN", "weight": 4}, {"word": "zarządzanie serwerami", "weight": 5},
            {"word": "backup", "weight": 4}, {"word": "monitoring sieci", "weight": 4},
            {"word": "rozwiązywanie problemów", "weight": 4}, {"word": "load balancing", "weight": 4},
            {"word": "NAT", "weight": 3}, {"word": "VLAN", "weight": 3},
            {"word": "QoS", "weight": 3}, {"word": "SNMP", "weight": 3},
            {"word": "Wi-Fi", "weight": 4}, {"word": "urządzenia sieciowe", "weight": 4},
            {"word": "Cisco", "weight": 5}, {"word": "Juniper", "weight": 4},
            {"word": "Mikrotik", "weight": 4}, {"word": "przełączniki", "weight": 4},
            {"word": "routery", "weight": 4}, {"word": "topologia sieci", "weight": 4},
            {"word": "systemy redundantne", "weight": 5}, {"word": "HA (High Availability)", "weight": 5},
            {"word": "IPSec", "weight": 4}, {"word": "Ethernet", "weight": 4},
            {"word": "802.11", "weight": 3}, {"word": "MPLS", "weight": 3},
            {"word": "IP", "weight": 4}, {"word": "IPv6", "weight": 4},
            {"word": "zarządzanie dostępem", "weight": 5}, {"word": "kontrola przepływu", "weight": 4},
            {"word": "proxy", "weight": 3}, {"word": "IDS/IPS", "weight": 4},
            {"word": "SIEM", "weight": 4}, {"word": "zarządzanie urządzeniami", "weight": 4},
            {"word": "PowerShell", "weight": 4}, {"word": "Bash", "weight": 4},
            {"word": "automatyzacja", "weight": 5}, {"word": "monitorowanie ruchu", "weight": 4}, 
            {"word": "B1", "weight": 1}, {"word": "B2", "weight": 2},
            {"word": "C1", "weight": 3}, {"word": "C2", "weight": 4}, {"word": "English", "weight": 1}, 
            {"word": "Angielski", "weight": 1}
        ],
    },
    {
        "title": "Specjalista HR",
        "keywords": [
            {"word": "rekrutacja", "weight": 5}, {"word": "onboarding", "weight": 5},
            {"word": "rozwój pracowników", "weight": 4}, {"word": "rozmowy kwalifikacyjne", "weight": 4},
            {"word": "ocena pracowników", "weight": 4}, {"word": "szkolenia", "weight": 4},
            {"word": "employer branding", "weight": 5}, {"word": "zarządzanie talentami", "weight": 5},
            {"word": "benefity", "weight": 4}, {"word": "kultura organizacyjna", "weight": 4},
            {"word": "motywacja", "weight": 4}, {"word": "zarządzanie konfliktami", "weight": 4},
            {"word": "HRIS", "weight": 4}, {"word": "umowy", "weight": 3},
            {"word": "rozwój zawodowy", "weight": 4}, {"word": "rotacja pracowników", "weight": 4},
            {"word": "strategia HR", "weight": 5}, {"word": "planowanie zatrudnienia", "weight": 5},
            {"word": "headhunting", "weight": 4}, {"word": "rekrutacja IT", "weight": 4},
            {"word": "analiza kompetencji", "weight": 4}, {"word": "zaangażowanie pracowników", "weight": 4},
            {"word": "raporty HR", "weight": 4}, {"word": "wskaźniki KPI", "weight": 4},
            {"word": "prawo pracy", "weight": 4}, {"word": "wynagrodzenie", "weight": 3},
            {"word": "wyniki", "weight": 3}, {"word": "programy motywacyjne", "weight": 4},
            {"word": "feedback", "weight": 4}, {"word": "rozwój liderów", "weight": 5},
            {"word": "harmonogramy", "weight": 3}, {"word": "polityka firmy", "weight": 3}, 
            {"word": "B1", "weight": 1}, {"word": "B2", "weight": 2},
            {"word": "C1", "weight": 3}, {"word": "C2", "weight": 4}, {"word": "English", "weight": 1}, 
            {"word": "Angielski", "weight": 1}
            
        ],
    },
    {
        "title": "Księgowy",
        "keywords": [
            {"word": "rachunkowość", "weight": 5}, {"word": "księgi rachunkowe", "weight": 4},
            {"word": "podatki", "weight": 5}, {"word": "VAT", "weight": 4},
            {"word": "PIT", "weight": 4}, {"word": "CIT", "weight": 4},
            {"word": "bilans", "weight": 4}, {"word": "sprawozdania finansowe", "weight": 4},
            {"word": "budżetowanie", "weight": 5}, {"word": "kontrola kosztów", "weight": 4},
            {"word": "SAP", "weight": 4}, {"word": "audyt", "weight": 4},
            {"word": "Excel", "weight": 4}, {"word": "raportowanie", "weight": 4},
            {"word": "analiza kosztów", "weight": 4}, {"word": "rozliczenia", "weight": 4},
            {"word": "fakturowanie", "weight": 4}, {"word": "amortyzacja", "weight": 3},
            {"word": "prognozowanie finansowe", "weight": 4}, {"word": "zgodność podatkowa", "weight": 4},
            {"word": "kontrole wewnętrzne", "weight": 4}, {"word": "analiza wydatków", "weight": 4},
            {"word": "bilansowanie", "weight": 4}, {"word": "systemy ERP", "weight": 4},
            {"word": "Cash Flow", "weight": 5}, {"word": "zobowiązania", "weight": 3},
            {"word": "należności", "weight": 3}, {"word": "księgowość zarządcza", "weight": 4},
            {"word": "analiza finansowa", "weight": 5}, {"word": "zgodność regulacyjna", "weight": 5},
            {"word": "optymalizacja podatkowa", "weight": 4}, {"word": "VAT-UE", "weight": 4},
            {"word": "koszt jednostkowy", "weight": 3}, {"word": "B1", "weight": 1}, {"word": "B2", "weight": 2},
            {"word": "C1", "weight": 3}, {"word": "C2", "weight": 4}, {"word": "English", "weight": 1}, 
            {"word": "Angielski", "weight": 1}
        ],
    },
    {
        "title": "Inżynier Mechanik",
        "keywords": [
            {"word": "projektowanie", "weight": 5}, {"word": "CAD", "weight": 5},
            {"word": "SolidWorks", "weight": 5}, {"word": "inżynieria", "weight": 4},
            {"word": "mechanika", "weight": 4}, {"word": "materiały", "weight": 4},
            {"word": "modelowanie 3D", "weight": 5}, {"word": "rysunki techniczne", "weight": 4},
            {"word": "analiza FEM", "weight": 5}, {"word": "automatyka", "weight": 4},
            {"word": "testowanie", "weight": 4}, {"word": "konserwacja", "weight": 3},
            {"word": "prototypowanie", "weight": 4}, {"word": "systemy mechaniczne", "weight": 5},
            {"word": "tolerancje", "weight": 4}, {"word": "termodynamika", "weight": 4},
            {"word": "dynamika", "weight": 4}, {"word": "projektowanie urządzeń", "weight": 5},
            {"word": "analiza zmęczeniowa", "weight": 5}, {"word": "maszyny przemysłowe", "weight": 4},
            {"word": "technologia produkcji", "weight": 4}, {"word": "kinematyka", "weight": 4},
            {"word": "statyka", "weight": 4}, {"word": "inżynieria procesowa", "weight": 5},
            {"word": "systemy hydrauliczne", "weight": 4}, {"word": "spawalnictwo", "weight": 3},
            {"word": "mechanika płynów", "weight": 4}, {"word": "pompy", "weight": 3},
            {"word": "przekładnie", "weight": 4}, {"word": "analiza naprężeń", "weight": 5},
            {"word": "symulacja komputerowa", "weight": 5}, {"word": "B1", "weight": 1}, {"word": "B2", "weight": 2},
            {"word": "C1", "weight": 3}, {"word": "C2", "weight": 4}
        ],
    },
    {
        "title": "Inżynier Budownictwa",
        "keywords": [
            {"word": "projektowanie", "weight": 5}, {"word": "AutoCAD", "weight": 5},
            {"word": "Revit", "weight": 5}, {"word": "budowa", "weight": 4},
            {"word": "kosztorysy", "weight": 4}, {"word": "struktury", "weight": 4},
            {"word": "inżynieria lądowa", "weight": 5}, {"word": "analiza", "weight": 4},
            {"word": "architektura", "weight": 4}, {"word": "nadzór", "weight": 4},
            {"word": "planowanie", "weight": 4}, {"word": "BIM", "weight": 5},
            {"word": "projekty budowlane", "weight": 5}, {"word": "konstrukcje stalowe", "weight": 5},
            {"word": "betony", "weight": 3}, {"word": "rysunki architektoniczne", "weight": 4},
            {"word": "geotechnika", "weight": 4}, {"word": "instalacje budowlane", "weight": 4},
            {"word": "zarządzanie budową", "weight": 5}, {"word": "dokumentacja techniczna", "weight": 4},
            {"word": "wytrzymałość materiałów", "weight": 5}, {"word": "pomiary", "weight": 4},
            {"word": "mosty", "weight": 5}, {"word": "drogi", "weight": 5},
            {"word": "budynki wielkopowierzchniowe", "weight": 5}, {"word": "budynki mieszkalne", "weight": 5},
            {"word": "inspekcje budowlane", "weight": 4}, {"word": "B1", "weight": 1}, {"word": "B2", "weight": 2},
            {"word": "C1", "weight": 3}, {"word": "C2", "weight": 4}
        ],
    },
    {
        "title": "Kierownik Sprzedaży",
        "keywords": [
            {"word": "sprzedaż", "weight": 5}, {"word": "zarządzanie klientami", "weight": 5},
            {"word": "KPI", "weight": 5}, {"word": "negocjacje", "weight": 5},
            {"word": "CRM", "weight": 5}, {"word": "budowanie relacji", "weight": 5},
            {"word": "lead generation", "weight": 5}, {"word": "strategia sprzedaży", "weight": 5},
            {"word": "targety", "weight": 4}, {"word": "raportowanie", "weight": 4},
            {"word": "prezentacje", "weight": 4}, {"word": "prognozy sprzedaży", "weight": 4},
            {"word": "analizy", "weight": 4}, {"word": "pipeline sprzedaży", "weight": 5},
            {"word": "pozyskiwanie klientów", "weight": 5}, {"word": "zarządzanie zespołem", "weight": 5},
            {"word": "follow-up", "weight": 4}, {"word": "utrzymanie klienta", "weight": 5},
            {"word": "marketing sprzedażowy", "weight": 4}, {"word": "współpraca z partnerami", "weight": 4},
            {"word": "networking", "weight": 4}, {"word": "B1", "weight": 1}, {"word": "B2", "weight": 2},
            {"word": "C1", "weight": 3}, {"word": "C2", "weight": 4}, {"word": "English", "weight": 1}, 
            {"word": "Angielski", "weight": 1}
        ],
    },
    {
        "title": "DevOps Engineer",
        "keywords": [
            {"word": "CI/CD", "weight": 5}, {"word": "Docker", "weight": 5},
            {"word": "Kubernetes", "weight": 5}, {"word": "Jenkins", "weight": 4},
            {"word": "automatyzacja", "weight": 5}, {"word": "monitoring", "weight": 4},
            {"word": "infrastruktura", "weight": 5}, {"word": "AWS", "weight": 5},
            {"word": "Azure", "weight": 5}, {"word": "chmura", "weight": 4},
            {"word": "skrypty", "weight": 5}, {"word": "Ansible", "weight": 5},
            {"word": "Terraform", "weight": 5}, {"word": "OpenShift", "weight": 4},
            {"word": "bezpieczeństwo DevOps", "weight": 4}, {"word": "zarządzanie logami", "weight": 4},
            {"word": "systemy kontenerowe", "weight": 5}, {"word": "zarządzanie wersjami", "weight": 4},
            {"word": "promowanie kodu", "weight": 4}, {"word": "load balancing", "weight": 4},
            {"word": "TDD", "weight": 5}, {"word": "GitOps", "weight": 5},
            {"word": "zarządzanie wdrożeniami", "weight": 5}, {"word": "pipeline", "weight": 4},
            {"word": "observability", "weight": 4}, {"word": "B1", "weight": 1}, {"word": "B2", "weight": 2},
            {"word": "C1", "weight": 3}, {"word": "C2", "weight": 4}, {"word": "English", "weight": 1}, 
            {"word": "Angielski", "weight": 1}
        ],
    },
    {
        "title": "Nauczyciel",
        "keywords": [
            {"word": "nauczanie", "weight": 5}, {"word": "program nauczania", "weight": 4},
            {"word": "lekcje", "weight": 4}, {"word": "pedagogika", "weight": 5},
            {"word": "materiały", "weight": 4}, {"word": "motywacja", "weight": 4},
            {"word": "oceny", "weight": 4}, {"word": "edukacja", "weight": 5},
            {"word": "testy", "weight": 4}, {"word": "egzaminy", "weight": 4},
            {"word": "dydaktyka", "weight": 5}, {"word": "metody nauczania", "weight": 4},
            {"word": "e-learning", "weight": 5}, {"word": "interakcja z uczniami", "weight": 4},
            {"word": "zarządzanie klasą", "weight": 4}, {"word": "innowacje edukacyjne", "weight": 4},
            {"word": "technologia edukacyjna", "weight": 4}, {"word": "B1", "weight": 1}, {"word": "B2", "weight": 2},
            {"word": "C1", "weight": 3}, {"word": "C2", "weight": 4}, {"word": "English", "weight": 1}, 
            {"word": "Angielski", "weight": 1}
        ],
    },
    {
        "title": "Specjalista Wsparcia IT",
        "keywords": [
            {"word": "wsparcie", "weight": 5}, {"word": "helpdesk", "weight": 5},
            {"word": "systemy operacyjne", "weight": 4}, {"word": "rozwiązywanie problemów", "weight": 5},
            {"word": "hardware", "weight": 4}, {"word": "software", "weight": 4},
            {"word": "serwis", "weight": 4}, {"word": "konfiguracja", "weight": 5},
            {"word": "zarządzanie urządzeniami", "weight": 4}, {"word": "dokumentacja", "weight": 4},
            {"word": "usuwanie awarii", "weight": 5}, {"word": "aktualizacje systemowe", "weight": 4},
            {"word": "instalacja oprogramowania", "weight": 5}, {"word": "ITIL", "weight": 5},
            {"word": "narzędzia wsparcia", "weight": 4}, {"word": "telefoniczne wsparcie techniczne", "weight": 4},
            {"word": "zarządzanie użytkownikami", "weight": 4}, {"word": "B1", "weight": 1}, {"word": "B2", "weight": 2},
            {"word": "C1", "weight": 3}, {"word": "C2", "weight": 4}, {"word": "English", "weight": 1}, 
            {"word": "Angielski", "weight": 1}
        ],
    }
    
    ]

    for pos in default_positions:
        # Sprawdź, czy stanowisko już istnieje
        if not Position.query.filter_by(title=pos["title"], is_default=True).first():
            position = Position(title=pos["title"], is_default=True)
            db.session.add(position)
            db.session.commit()

            # Dodaj słowa kluczowe z wagami
            for keyword in pos["keywords"]:
                kw = Keyword(
                    word=keyword["word"],
                    weight=keyword["weight"],
                    position_id=position.id,
                )
                db.session.add(kw)

    db.session.commit()



def create_global_positions():
    from models import Position, Keyword

    global_positions = [
        {
            "title": "Programista",
            "keywords": [
                "programowanie",
                "Python",
                "Java",
                "JavaScript",
                "C#",
                "SQL",
                "bazy danych",
                "API",
                "framework",
                "debugowanie",
                "kodowanie",
                "systemy operacyjne",
                "integracja",
                "automatyzacja",
                "scrum",
                "agile",
                "testing",
                "backend",
                "frontend",
                "narzędzia CI/CD",
                "Git",
                "Docker",
                "Linux",
                "Windows",
                "architektura systemów",
                "analityka",
                "optymalizacja",
                "REST",
                "SOAP",
                "mikroserwisy",
                "bezpieczeństwo IT",
                "refaktoryzacja",
                "chmura",
                "Azure",
                "AWS",
                "GCP",
                "DevOps",
            ],
        },
        {
            "title": "Analityk Danych",
            "keywords": [
                "dane",
                "Excel",
                "SQL",
                "analizy",
                "wizualizacja",
                "Tableau",
                "Power BI",
                "statystyka",
                "Python",
                "R",
                "modele predykcyjne",
                "algorytmy",
                "analiza danych",
                "przetwarzanie danych",
                "ETL",
                "dashboard",
                "raporty",
                "Big Data",
                "Hadoop",
                "Spark",
                "analityka biznesowa",
                "insighty",
                "dane sprzedażowe",
                "prognozowanie",
                "optymalizacja",
                "bazy danych",
                "machine learning",
                "analiza rynku",
                "czyszczenie danych",
                "Google Analytics",
                "KPIs",
                "modelowanie danych",
            ],
        },
        {
            "title": "Kierownik Projektu",
            "keywords": [
                "zarządzanie projektami",
                "harmonogramy",
                "budżet",
                "agile",
                "scrum",
                "waterfall",
                "zarządzanie ryzykiem",
                "priorytetyzacja",
                "planowanie",
                "komunikacja",
                "zarządzanie zespołem",
                "stakeholderzy",
                "analiza wymagań",
                "cele projektu",
                "raportowanie",
                "monitoring",
                "zarządzanie czasem",
                "zarządzanie zmianą",
                "leadership",
                "efektywność",
                "metodologia",
                "stand-up",
                "roadmap",
                "narzędzia PM",
                "Jira",
                "Asana",
                "MS Project",
                "zarządzanie zadaniami",
            ],
        },
        {
            "title": "Grafik",
            "keywords": [
                "projektowanie graficzne",
                "Adobe Photoshop",
                "Illustrator",
                "InDesign",
                "Canva",
                "projektowanie logo",
                "typografia",
                "paleta kolorów",
                "branding",
                "layout",
                "infografiki",
                "kreatywność",
                "ilustracje",
                "UX/UI",
                "projektowanie interfejsów",
                "responsywność",
                "druk",
                "składy",
                "grafika wektorowa",
                "animacje",
                "mockupy",
                "wizualizacje",
                "wireframes",
                "projektowanie banerów",
                "grafika 3D",
                "narzędzia graficzne",
            ],
        },
        {
            "title": "Copywriter",
            "keywords": [
                "SEO",
                "optymalizacja treści",
                "copywriting",
                "marketing treści",
                "artykuły",
                "blogi",
                "storytelling",
                "kreatywne pisanie",
                "redagowanie",
                "social media",
                "CTA",
                "reklama",
                "edytowanie",
                "język",
                "treść",
                "analiza konkurencji",
                "Google Analytics",
                "targetowanie",
                "kampanie marketingowe",
                "e-mail marketing",
                "UX writing",
                "optymalizacja słów kluczowych",
                "analiza odbiorców",
            ],
        },
        {
            "title": "Specjalista ds. Marketingu Cyfrowego",
            "keywords": [
                "marketing cyfrowy",
                "SEO",
                "SEM",
                "Google Ads",
                "Facebook Ads",
                "social media",
                "kampanie",
                "reklama",
                "targetowanie",
                "optymalizacja",
                "Google Analytics",
                "email marketing",
                "automatyzacja marketingu",
                "content marketing",
                "remarketing",
                "KPI",
                "lead generation",
                "konwersje",
                "budżet reklamowy",
                "strategia marketingowa",
            ],
        },
        {
            "title": "Administrator Sieci",
            "keywords": [
                "sieci komputerowe",
                "protokoły",
                "TCP/IP",
                "DNS",
                "DHCP",
                "routing",
                "VPN",
                "bezpieczeństwo",
                "firewalle",
                "zarządzanie siecią",
                "konfiguracja",
                "LAN",
                "WAN",
                "zarządzanie serwerami",
                "backup",
                "monitoring sieci",
                "rozwiązywanie problemów",
                "load balancing",
                "NAT",
                "VLAN",
                "QoS",
                "SNMP",
                "Wi-Fi",
                "urządzenia sieciowe",
                "Cisco",
                "Juniper",
                "Mikrotik",
                "przełączniki",
                "routery",
                "topologia sieci",
                "systemy redundantne",
                "HA (High Availability)",
                "IPSec",
                "Ethernet",
                "802.11",
                "MPLS",
                "IP",
                "IPv6",
                "zarządzanie dostępem",
                "kontrola przepływu",
                "proxy",
                "IDS/IPS",
                "SIEM",
                "zarządzanie urządzeniami",
                "PowerShell",
                "Bash",
                "automatyzacja",
                "monitorowanie ruchu",
            ],
        },
        {
            "title": "Specjalista HR",
            "keywords": [
                "rekrutacja",
                "onboarding",
                "rozwój pracowników",
                "rozmowy kwalifikacyjne",
                "ocena pracowników",
                "szkolenia",
                "employer branding",
                "zarządzanie talentami",
                "benefity",
                "kultura organizacyjna",
                "motywacja",
                "zarządzanie konfliktami",
                "HRIS",
                "umowy",
                "rozwój zawodowy",
                "rotacja pracowników",
                "strategia HR",
                "planowanie zatrudnienia",
                "headhunting",
                "rekrutacja IT",
                "analiza kompetencji",
                "zaangażowanie pracowników",
                "raporty HR",
                "wskaźniki KPI",
                "prawo pracy",
                "wynagrodzenie",
                "wyniki",
                "programy motywacyjne",
                "feedback",
                "rozwój liderów",
                "harmonogramy",
                "polityka firmy",
            ],
        },
        {
            "title": "Księgowy",
            "keywords": [
                "rachunkowość",
                "księgi rachunkowe",
                "podatki",
                "VAT",
                "PIT",
                "CIT",
                "bilans",
                "sprawozdania finansowe",
                "budżetowanie",
                "kontrola kosztów",
                "SAP",
                "audyt",
                "Excel",
                "raportowanie",
                "analiza kosztów",
                "rozliczenia",
                "fakturowanie",
                "amortyzacja",
                "prognozowanie finansowe",
                "zgodność podatkowa",
                "kontrole wewnętrzne",
                "analiza wydatków",
                "bilansowanie",
                "systemy ERP",
                "Cash Flow",
                "zobowiązania",
                "należności",
                "księgowość zarządcza",
                "analiza finansowa",
                "zgodność regulacyjna",
                "optymalizacja podatkowa",
                "VAT-UE",
                "koszt jednostkowy",
            ],
        },
        {
            "title": "Inżynier Mechanik",
            "keywords": [
                "projektowanie",
                "CAD",
                "SolidWorks",
                "inżynieria",
                "mechanika",
                "materiały",
                "modelowanie 3D",
                "rysunki techniczne",
                "analiza FEM",
                "automatyka",
                "testowanie",
                "konserwacja",
                "prototypowanie",
                "systemy mechaniczne",
                "tolerancje",
                "termodynamika",
                "dynamika",
                "projektowanie urządzeń",
                "analiza zmęczeniowa",
                "maszyny przemysłowe",
                "technologia produkcji",
                "kinematyka",
                "statyka",
                "inżynieria procesowa",
                "systemy hydrauliczne",
                "spawalnictwo",
                "mechanika płynów",
                "pompy",
                "przekładnie",
                "analiza naprężeń",
                "symulacja komputerowa",
            ],
        },
        {
            "title": "Inżynier Budownictwa",
            "keywords": [
                "projektowanie",
                "AutoCAD",
                "Revit",
                "budowa",
                "kosztorysy",
                "struktury",
                "inżynieria lądowa",
                "analiza",
                "architektura",
                "nadzór",
                "planowanie",
                "BIM",
                "projekty budowlane",
                "konstrukcje stalowe",
                "betony",
                "rysunki architektoniczne",
                "geotechnika",
                "instalacje budowlane",
                "zarządzanie budową",
                "dokumentacja techniczna",
                "wytrzymałość materiałów",
                "pomiary",
                "mosty",
                "drogi",
                "budynki wielkopowierzchniowe",
                "budynki mieszkalne",
                "inspekcje budowlane",
            ],
        },
        {
            "title": "Kierownik Sprzedaży",
            "keywords": [
                "sprzedaż",
                "zarządzanie klientami",
                "KPI",
                "negocjacje",
                "CRM",
                "budowanie relacji",
                "lead generation",
                "strategia sprzedaży",
                "targety",
                "raportowanie",
                "prezentacje",
                "prognozy sprzedaży",
                "analizy",
                "pipeline sprzedaży",
                "pozyskiwanie klientów",
                "zarządzanie zespołem",
                "follow-up",
                "utrzymanie klienta",
                "marketing sprzedażowy",
                "współpraca z partnerami",
                "networking",
            ],
        },
        {
            "title": "DevOps Engineer",
            "keywords": [
                "CI/CD",
                "Docker",
                "Kubernetes",
                "Jenkins",
                "automatyzacja",
                "monitoring",
                "infrastruktura",
                "AWS",
                "Azure",
                "chmura",
                "skrypty",
                "Ansible",
                "Terraform",
                "OpenShift",
                "bezpieczeństwo DevOps",
                "zarządzanie logami",
                "systemy kontenerowe",
                "zarządzanie wersjami",
                "promowanie kodu",
                "load balancing",
                "TDD",
                "GitOps",
                "zarządzanie wdrożeniami",
                "pipeline",
                "observability",
            ],
        },
        {
            "title": "Nauczyciel",
            "keywords": [
                "nauczanie",
                "program nauczania",
                "lekcje",
                "pedagogika",
                "materiały",
                "motywacja",
                "oceny",
                "edukacja",
                "testy",
                "egzaminy",
                "dydaktyka",
                "metody nauczania",
                "e-learning",
                "interakcja z uczniami",
                "zarządzanie klasą",
                "innowacje edukacyjne",
                "technologia edukacyjna",
            ],
        },
        {
            "title": "Specjalista Wsparcia IT",
            "keywords": [
                "wsparcie",
                "helpdesk",
                "systemy operacyjne",
                "rozwiązywanie problemów",
                "hardware",
                "software",
                "serwis",
                "konfiguracja",
                "zarządzanie urządzeniami",
                "dokumentacja",
                "usuwanie awarii",
                "aktualizacje systemowe",
                "instalacja oprogramowania",
                "ITIL",
                "narzędzia wsparcia",
                "telefoniczne wsparcie techniczne",
                "zarządzanie użytkownikami",
            ],
        },
    ]

    for position_data in global_positions:
        existing_position = Position.query.filter_by(
            title=position_data["title"], is_global=True
        ).first()
        if not existing_position:
            new_position = Position(title=position_data["title"], is_global=True)
            db.session.add(new_position)
            db.session.commit()
            for keyword in position_data["keywords"]:
                new_keyword = Keyword(word=keyword, position_id=new_position.id)
                db.session.add(new_keyword)
            db.session.commit()


# Uruchamianie aplikacji
if __name__ == "__main__":
    app = create_app()
    app.run(debug=True)
